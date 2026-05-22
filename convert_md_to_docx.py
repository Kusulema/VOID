#!/usr/bin/env python3
"""
Простой конвертер Markdown -> DOCX для локального использования.
Этот скрипт не поддерживает полноценный Markdown -> DOCX (используйте Pandoc для лучшего результата),
но создаёт .docx с заголовками и параграфами и сохраняет маркеры `[[screenshot:...]]`.

Установка зависимостей:
    pip install python-docx

Запуск:
    python convert_md_to_docx.py Lopputoo_final.md Lõpputöö4_filled.docx

Если Pandoc установлен, предпочтительнее запустить:
    pandoc Lopputoo_final.md -o Lõpputöö4_filled.docx

"""
import sys
from docx import Document
from docx.shared import Pt

def convert(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.rstrip('\n')
            if not line:
                doc.add_paragraph('')
                continue
            # заголовки
            if line.startswith('### '):
                p = doc.add_heading(line[4:].strip(), level=3)
                continue
            if line.startswith('## '):
                p = doc.add_heading(line[3:].strip(), level=2)
                continue
            if line.startswith('# '):
                p = doc.add_heading(line[2:].strip(), level=1)
                continue
            # маркеры скриншотов оставляем как отдельный параграф с пометкой
            if line.startswith('[[screenshot:') and line.endswith(']]'):
                desc = line[len('[[screenshot:'):-2]
                p = doc.add_paragraph('[[SCREENSHOT: {} ]]'.format(desc))
                p.runs[0].font.italic = True
                continue
            # простой маркер кода или выделения сохраняем как есть
            doc.add_paragraph(line)

    doc.save(docx_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: convert_md_to_docx.py source.md output.docx')
        sys.exit(2)
    convert(sys.argv[1], sys.argv[2])
